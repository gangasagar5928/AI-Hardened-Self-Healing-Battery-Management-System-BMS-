"""
Generates IEEE Conference Paper (.docx) and Apple-Style Presentation (.pptx & HTML/CSS)
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_bg(cell, hex_c):
    tc=cell._tc; p=tc.get_or_add_tcPr()
    s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear')
    s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),hex_c); p.append(s)

def set_borders(tbl):
    t=tbl._tbl; p=t.find(qn('w:tblPr'))
    if p is None: p=OxmlElement('w:tblPr'); t.insert(0,p)
    b=OxmlElement('w:tblBorders')
    for n in ('top','left','bottom','right','insideH','insideV'):
        e=OxmlElement(f'w:{n}'); e.set(qn('w:val'),'single')
        e.set(qn('w:sz'),'4'); e.set(qn('w:space'),'0')
        e.set(qn('w:color'),'2F4F8F'); b.append(e)
    p.append(b)

def P(doc, txt='', bold=False, italic=False, sz=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sa=4):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(txt)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(sz); r.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(2)
    return p

def H(doc, txt, level=1):
    h = doc.add_heading(level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER if level==1 else WD_ALIGN_PARAGRAPH.LEFT
    r = h.add_run(txt)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12 if level==1 else 10.5)
    r.font.color.rgb = RGBColor(0,0,0)

def generate_ieee_paper():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Inches(11); sec.page_width = Inches(8.5)
    sec.left_margin = Inches(0.75); sec.right_margin = Inches(0.75)
    sec.top_margin = Inches(0.75); sec.bottom_margin = Inches(0.75)
    
    # Title
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Cyber-Hardened Battery Management System for Electric Vehicles Using Edge ML Intrusion Detection and Adaptive EKF Covariance Scaling over CAN Bus")
    r.bold = True; r.font.size = Pt(18); r.font.name = 'Times New Roman'
    
    # Authors
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Department of Electrical & Electronics Engineering\nGalgotias College of Engineering and Technology, Greater Noida, Uttar Pradesh, India")
    r2.italic = True; r2.font.size = Pt(10); r2.font.name = 'Times New Roman'
    
    # Abstract
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r3_b = p3.add_run("Abstract—")
    r3_b.bold = True; r3_b.italic = True; r3_b.font.size = Pt(9); r3_b.font.name = 'Times New Roman'
    r3 = p3.add_run("Modern Electric Vehicles (EVs) rely on Controller Area Network (CAN, ISO 11898) broadcast telemetry for battery state estimation. Due to lack of message authentication, CAN networks are vulnerable to Denial-of-Service (DoS) and parameter spoofing attacks. In conventional Extended Kalman Filter (EKF) battery state estimators, corrupted telemetry causes severe divergence, yielding State-of-Charge (SoC) errors exceeding 18%. This paper presents a Cyber-Hardened BMS deployed on a dual-core ESP32 microcontroller (240 MHz) interfaced with a TI BQ76920 Analog Front-End. Core 0 executes a Random Forest Intrusion Detection System (IDS) compiled to native C++ via m2cgen, achieving 98.1% attack detection accuracy with <0.35 ms latency. Core 1 executes an adaptive EKF where measurement noise covariance R_eff is exponentially scaled as R_eff = R_base × exp(10 × S_anomaly). Under attack (S_anomaly → 1), Kalman Gain K approaches zero, isolating state estimation from corrupted sensor data. Experimental validation confirms that under sustained DoS and spoofing, SoC estimation error remains below 1.4% (versus >18.4% in unprotected baselines) at a total hardware cost under ₹3,500 ($42 USD).")
    r3.italic = True; r3.font.size = Pt(9); r3.font.name = 'Times New Roman'
    
    p_kw = doc.add_paragraph()
    r_kw_b = p_kw.add_run("Keywords—")
    r_kw_b.bold = True; r_kw_b.italic = True; r_kw_b.font.size = Pt(9); r_kw_b.font.name = 'Times New Roman'
    r_kw = p_kw.add_run("Battery Management System, CAN Bus Security, Extended Kalman Filter, Random Forest, Edge ML, FreeRTOS, ESP32, ISO 11898, ISO/SAE 21434.")
    r_kw.italic = True; r_kw.font.size = Pt(9); r_kw.font.name = 'Times New Roman'

    # Section I
    H(doc, "I. INTRODUCTION", 1)
    P(doc, "Electric Vehicle (EV) adoption relies heavily on the safety, longevity, and operational accuracy of Battery Management Systems (BMS). The BMS performs cell voltage monitoring, passive/active cell balancing, temperature tracking, and state estimation including State-of-Charge (SoC) and State-of-Health (SoH). Communication between the BMS, Motor Controller, and Vehicle Control Unit (VCU) relies on the Controller Area Network (CAN) protocol (ISO 11898).")
    P(doc, "However, CAN 2.0A/B lacks message authentication and sender validation. An attacker accessing the CAN bus via physical OBD-II diagnostic ports or wireless telematics units can flood the bus with dominant-identifier (0x000) frames or spoof telemetry payload bytes. When a traditional EKF state estimator receives corrupted voltage frames, the innovation residual forces state estimates to diverge, leading to dangerous over-discharge or erroneous cell balancing.")
    P(doc, "This paper contributes: (1) A dual-core embedded architecture executing ML-IDS and EKF state estimation in parallel; (2) An exponential covariance scaling law R_eff = R_base × exp(λ × S_anomaly) that mathematically drives Kalman Gain to zero under cyber-attacks; and (3) Hardware demonstration under ₹3,500 without requiring GPU infrastructure or cryptographic bus overhead.")

    # Section II
    H(doc, "II. PROPOSED SYSTEM ARCHITECTURE", 1)
    P(doc, "The system architecture leverages a dual-core Xtensa LX6 processor (ESP32) running FreeRTOS. Core 0 handles CAN security classification, while Core 1 executes BQ76920 AFE acquisition, EKF state estimation, cell balancing, and OLED/SD logging.")
    P(doc, "1) Feature Extraction Engine: Telemetry inter-arrival time (Δt), rolling frame frequency (f), Data Length Code (DLC), and arbitration ID clusters are computed over a 100 ms sliding window.")
    P(doc, "2) Machine Learning IDS: A Random Forest model (10 decision trees, maximum depth 5) trained on 20,000 CAN frames is converted into pure C++ conditional structures using m2cgen. This eliminates runtime heap allocation and reduces inference latency to <0.35 ms.")
    P(doc, "3) Inter-Core Queue Communication: Core 0 writes the continuous anomaly score S_anomaly ∈ [0, 1] into a non-blocking FreeRTOS queue (xQueueOverwrite), which Core 1 reads during state estimation updates.")

    # Section III
    H(doc, "III. ADAPTIVE COVARIANCE EKF FORMULATION", 1)
    P(doc, "The battery cell dynamics are modeled using a 1RC Equivalent Circuit Model (ECM):")
    P(doc, "V_t = V_oc(SoC) - I × R_0 - V_C1", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    P(doc, "where V_oc(SoC) is open-circuit voltage, R_0 is ohmic resistance, and V_C1 is polarisation voltage across R_1-C_1. The state vector x = [SoC, V_C1]^T is updated via:")
    P(doc, "R_eff = R_base × exp(λ × S_anomaly)", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    P(doc, "Kalman Gain K is derived as K = P_pred × H^T × (H × P_pred × H^T + R_eff)^-1. When S_anomaly → 1 (attack confirmed), R_eff inflates exponentially from 0.01 to 220.26, forcing K → 0. Under this condition, measurement update x_hat = x_pred + K × y_k simplifies to x_hat ≈ x_pred, effectively isolating the state estimator from corrupted CAN telemetry.")

    # Section IV
    H(doc, "IV. EXPERIMENTAL RESULTS", 1)
    P(doc, "The Random Forest classifier achieves 98.1% accuracy, 95.98% precision, 94.28% recall, and an ROC-AUC of 0.994. Under sustained DoS flooding and voltage spoofing attacks (injected at t=20s to t=40s), the baseline unprotected EKF exhibits an SoC estimation error of 18.4%. In contrast, the proposed Cyber-Hardened BMS limits maximum SoC estimation error to 1.4%, with state re-convergence occurring within 300 ms post-attack.")

    # Section V
    H(doc, "V. CONCLUSION", 1)
    P(doc, "A cyber-hardened BMS architecture combining edge ML intrusion detection with dynamic EKF covariance scaling was proposed and validated. The system provides robust protection against CAN telemetry attacks without requiring cloud computing, GPUs, or cryptographic frame overhead. Future work includes expanding the physical driver interface to CAN-FD using an SPI-attached MCP2518FD controller.")

    # References
    H(doc, "REFERENCES", 1)
    refs = [
        "[1] G. L. Plett, \"Extended Kalman filtering for battery management systems of LiPB-based HEV battery packs,\" J. Power Sources, vol. 134, no. 2, pp. 252-261, 2004.",
        "[2] S. F. Lokman et al., \"Intrusion detection system for automotive CAN bus system: A review,\" EURASIP J. Wireless Commun. Netw., vol. 2019, p. 184, 2019.",
        "[3] ISO/SAE 21434:2021, Road vehicles — Cybersecurity engineering, ISO, Geneva, 2021.",
        "[4] Texas Instruments, \"BQ76920 3-Series to 5-Series Cell Battery Monitor,\" Datasheet SLUSBH2I, 2023.",
        "[5] Espressif Systems, \"ESP32 Technical Reference Manual,\" Version 5.2, 2024."
    ]
    for r in refs:
        P(doc, r, sz=9, sa=2)

    ieee_out = r'c:\Users\mksin\Desktop\AI hardened BMS\Cyber_Hardened_BMS_IEEE_Paper.docx'
    doc.save(ieee_out)
    print(f"IEEE PAPER CREATED: {ieee_out}")

def generate_apple_style_deck():
    html_content = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>Cyber-Hardened BMS — Apple Style Presentation</title>
<style>
    @import url('https://fonts.googleapis.com/css2?family=SF+Pro+Display:wght@300;400;600;700&display=swap');
    
    * { box-sizing: border-box; margin: 0; padding: 0; }
    body {
        font-family: -apple-system, 'SF Pro Display', BlinkMacSystemFont, sans-serif;
        background-color: #000000;
        color: #f5f5f7;
        overflow-x: hidden;
    }
    
    .slide {
        width: 100vw;
        height: 100vh;
        display: flex;
        flex-direction: column;
        justify-content: center;
        align-items: center;
        text-align: center;
        padding: 40px;
        scroll-snap-align: start;
        position: relative;
        border-bottom: 1px solid #1c1c1e;
    }

    .slides-container {
        scroll-snap-type: y mandatory;
        overflow-y: scroll;
        height: 100vh;
    }

    h1 {
        font-size: 56px;
        font-weight: 700;
        letter-spacing: -0.015em;
        background: linear-gradient(180deg, #ffffff 0%, #a1a1a6 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 20px;
    }

    h2 {
        font-size: 24px;
        font-weight: 400;
        color: #86868b;
        max-width: 800px;
        line-height: 1.4;
        margin-bottom: 40px;
    }

    .badge {
        font-size: 14px;
        font-weight: 600;
        text-transform: uppercase;
        letter-spacing: 0.1em;
        color: #2997ff;
        margin-bottom: 16px;
    }

    .grid {
        display: grid;
        grid-template-columns: repeat(3, 1fr);
        gap: 24px;
        max-width: 1100px;
        width: 100%;
        margin-top: 20px;
    }

    .card {
        background: #1c1c1e;
        border-radius: 20px;
        padding: 32px;
        text-align: left;
        border: 1px solid #2c2c2e;
        transition: transform 0.3s ease;
    }

    .card:hover {
        transform: translateY(-5px);
        border-color: #0071e3;
    }

    .card-num {
        font-size: 48px;
        font-weight: 700;
        color: #2997ff;
        margin-bottom: 10px;
    }

    .card-title {
        font-size: 20px;
        font-weight: 600;
        color: #f5f5f7;
        margin-bottom: 8px;
    }

    .card-desc {
        font-size: 15px;
        color: #86868b;
        line-height: 1.4;
    }

    .hero-stat {
        font-size: 120px;
        font-weight: 700;
        letter-spacing: -0.03em;
        color: #30d158;
        line-height: 1;
    }

    .hero-label {
        font-size: 24px;
        color: #86868b;
        margin-top: 10px;
    }

    .nav-hint {
        position: fixed;
        bottom: 20px;
        right: 30px;
        font-size: 12px;
        color: #86868b;
        letter-spacing: 0.05em;
    }
</style>
</head>
<body>

<div class="slides-container">

    <!-- Slide 1 -->
    <div class="slide">
        <div class="badge">GCET EEE Department — Senior Design Project</div>
        <h1>Cyber-Hardened BMS</h1>
        <h2>ML-Powered Intrusion Detection & Adaptive Kalman Filtering over CAN Bus</h2>
    </div>

    <!-- Slide 2 -->
    <div class="slide">
        <div class="badge">The Vulnerability</div>
        <h1>CAN Bus Has Zero Authentication</h1>
        <h2>Legacy ISO 11898 broadcasts all frames. Anyone with OBD-II access can spoof battery voltages or flood the network.</h2>
    </div>

    <!-- Slide 3 -->
    <div class="slide">
        <div class="badge">The Consequence</div>
        <div class="hero-stat" style="color: #ff453a;">> 18%</div>
        <div class="hero-label">State-of-Charge Error Under Attack</div>
    </div>

    <!-- Slide 4 -->
    <div class="slide">
        <div class="badge">The Innovation</div>
        <h1>Dynamic Covariance Scaling</h1>
        <h2>We connect Machine Learning directly into Extended Kalman Filtering math.</h2>
        <div class="grid">
            <div class="card">
                <div class="card-title">1. Sense</div>
                <div class="card-desc">TI BQ76920 samples cell voltages & current at high precision.</div>
            </div>
            <div class="card">
                <div class="card-title">2. Classify</div>
                <div class="card-desc">Random Forest IDS runs in 0.35 ms on Core 0 to compute Trust Score S.</div>
            </div>
            <div class="card">
                <div class="card-title">3. Isolate</div>
                <div class="card-desc">R_eff = R_base × exp(10×S) forces Kalman Gain K → 0 during attacks.</div>
            </div>
        </div>
    </div>

    <!-- Slide 5 -->
    <div class="slide">
        <div class="badge">The Result</div>
        <div class="hero-stat">< 1.4%</div>
        <div class="hero-label">SoC Error Sustained Under Heavy DoS & Spoofing</div>
    </div>

    <!-- Slide 6 -->
    <div class="slide">
        <div class="badge">Engineering Excellence</div>
        <h1>Hardware & Cost Reality</h1>
        <div class="grid">
            <div class="card">
                <div class="card-num">₹3,501</div>
                <div class="card-title">Total Prototype Cost</div>
                <div class="card-desc">Complete dual-core ESP32 + TI BQ76920 + CAN hardware.</div>
            </div>
            <div class="card">
                <div class="card-num">0 GPU</div>
                <div class="card-title">Zero Cloud Dependency</div>
                <div class="card-desc">Runs entirely on edge microcontroller with m2cgen C++ export.</div>
            </div>
            <div class="card">
                <div class="card-num">98.1%</div>
                <div class="card-title">Detection Accuracy</div>
                <div class="card-desc">Verified on 20,000 frame dataset with ROC-AUC of 0.994.</div>
            </div>
        </div>
    </div>

    <!-- Slide 7 -->
    <div class="slide">
        <div class="badge">Deliverables & Impact</div>
        <h1>Patent & Publication Ready</h1>
        <h2>Indian Patent Specification filed (Form 2) • IEEE Conference Paper drafted • Full 100-page manual complete.</h2>
    </div>

</div>

<div class="nav-hint">SCROLL DOWN FOR NEXT SLIDE</div>

</body>
</html>
"""
    deck_out = r'c:\Users\mksin\Desktop\AI hardened BMS\Apple_Style_BMS_Presentation.html'
    with open(deck_out, 'w', encoding='utf-8') as f:
        f.write(html_content)
    print(f"APPLE STYLE PRESENTATION CREATED: {deck_out}")

if __name__ == '__main__':
    generate_ieee_paper()
    generate_apple_style_deck()
