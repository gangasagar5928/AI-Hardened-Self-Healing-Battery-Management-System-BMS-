import docx

doc_dl = docx.Document(r"C:\Users\mksin\Downloads\Cyber-Hardened-BMS-Complete-Manual.docx")

print("DOWNLOADS MANUAL STATS:")
print("Paragraphs:", len(doc_dl.paragraphs))
print("Tables:", len(doc_dl.tables))
words = sum(len(p.text.split()) for p in doc_dl.paragraphs)
for t in doc_dl.tables:
    for row in t.rows:
        for cell in row.cells:
            words += len(cell.text.split())
print("Total Words (Paragraphs + Tables):", words)

headings = []
for p in doc_dl.paragraphs:
    txt = p.text.strip()
    if txt.startswith("Chapter") or txt.startswith("APPENDIX") or txt.startswith("TABLE OF CONTENTS") or p.style.name.startswith("Heading"):
        headings.append(txt)

print("\n--- FIRST 40 HEADINGS IN DOWNLOADS FILE ---")
for h in headings[:40]:
    print(h)
