import shutil
import docx

src = r"C:\Users\mksin\Downloads\Cyber-Hardened-BMS-Complete-Manual.docx"
dst = r"c:\Users\mksin\Desktop\AI hardened BMS\Cyber_Hardened_BMS_Manual.docx"

# Copy source to destination
shutil.copy2(src, dst)

doc = docx.Document(dst)

# Count words
para_words = sum(len(p.text.split()) for p in doc.paragraphs)
table_words = 0
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            table_words += len(cell.text.split())

total_words = para_words + table_words

# Count images
img_count = sum(1 for rel in doc.part.rels.values() if 'image' in rel.target_ref)

print("="*60)
print("MASTER MANUAL VERIFICATION REPORT:")
print(f"File Location: {dst}")
print(f"Total Paragraphs: {len(doc.paragraphs)}")
print(f"Total Tables: {len(doc.tables)}")
print(f"Paragraph Words: {para_words}")
print(f"Table Words: {table_words}")
print(f"TOTAL WORD COUNT: {total_words}")
print(f"EMBEDDED IMAGES: {img_count}")
print("="*60)
