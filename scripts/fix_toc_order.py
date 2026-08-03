import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

file_path = r"c:\Users\mksin\Desktop\AI hardened BMS\Cyber_Hardened_BMS_Manual.docx"
doc = docx.Document(file_path)

# First remove any previously inserted TOC entries between TABLE OF CONTENTS and Chapter 1
toc_start_idx = -1
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "TABLE OF CONTENTS":
        toc_start_idx = i
        break

print(f"TABLE OF CONTENTS found at index: {toc_start_idx}")

# Delete paragraphs between TABLE OF CONTENTS and Chapter 1 body text
del_indices = []
for i in range(toc_start_idx + 1, len(doc.paragraphs)):
    txt = doc.paragraphs[i].text.strip()
    if txt.startswith("Chapter 1") and "Executive Summary" in txt and not "............" in txt:
        # Reached the actual Chapter 1 heading
        break
    else:
        del_indices.append(i)

# Delete in reverse order to keep indices valid
for idx in reversed(del_indices):
    p = doc.paragraphs[idx]
    p._element.getparent().remove(p._element)

# Re-fetch document after deletion
doc = docx.Document(file_path)
toc_start_idx = -1
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "TABLE OF CONTENTS":
        toc_start_idx = i
        break

toc_heading_p = doc.paragraphs[toc_start_idx]

# Define TOC entries in correct chronological order
toc_entries = [
    ("Chapter 1", "Executive Summary"),
    ("Chapter 2", "Introduction & Motivation"),
    ("Chapter 3", "Literature Review"),
    ("Chapter 4", "System Overview & Master Schematic Architecture"),
    ("Chapter 5", "Core Theory: How Everything Works"),
    ("Chapter 6", "Software & Tools Setup Guide"),
    ("Chapter 7", "Hardware Architecture & Bill of Materials"),
    ("Chapter 8", "Simulation Phase (LTspice & MATLAB/Simulink)"),
    ("Chapter 9", "Hardware Assembly (Safe, Step-by-Step)"),
    ("Chapter 10", "Firmware Development & Dual-Core Task Architecture"),
    ("Chapter 11", "Attack Bench & Dataset Generation"),
    ("Chapter 12", "ML Classifier Training & Deployment"),
    ("Chapter 13", "The IDS–EKF Feedback Loop (Patent Core)"),
    ("Chapter 14", "Testing, Validation & Deliverables Checklist"),
    ("Chapter 15", "12-Week Timeline & Milestones"),
    ("Chapter 16", "Writing the IEEE Conference Paper"),
    ("Chapter 17", "Patent Filing Guide (Indian Patent Office)"),
    ("Chapter 18", "Presenting to Your Professor & Viva Walkthrough"),
    ("Chapter 19", "Error Log: Every Correction Applied in This Manual"),
    ("Chapter 20", "Extended Kalman Filter: Full Mathematical Derivation"),
    ("Chapter 21", "Worked Numerical Examples & Step-by-Step Calculations"),
    ("Chapter 22", "Complete Firmware Source Code Listings"),
    ("Chapter 23", "Complete Python ML Pipeline Source Code"),
    ("Chapter 24", "Hardware & Software Troubleshooting Guide"),
    ("Chapter 25", "Frequently Asked Questions (FAQ)"),
    ("Chapter 26", "Team Roles & Daily Task Breakdown"),
    ("Chapter 27", "IEEE Paper: Full Draft Template"),
    ("Chapter 28", "Safety & Compliance Notes"),
    ("Chapter 29", "Component Datasheet Quick-Reference"),
    ("Chapter 30", "Standard Project Report Structure"),
    ("Chapter 31", "Patent Forms: Field-by-Field Guide"),
    ("Chapter 32", "Viva / Interview Questions & Model Answers"),
    ("Chapter 33", "Expected Results: Graph-by-Graph Description"),
    ("Chapter 34", "PCB Design Guide (KiCad Detail)"),
    ("Chapter 35", "Alternatives Considered and Rejected"),
    ("Chapter 36", "Data Logging & Post-Processing"),
    ("Chapter 37", "Environmental & Sustainability Notes"),
    ("Chapter 38", "Sample Data & Test Log Format"),
    ("Chapter 39", "Deliverables Mapped to Typical Evaluation Criteria"),
    ("Chapter 40", "Advanced Battery Diagnostics & Degradation Physics"),
    ("Chapter 41", "CAN Protocol Deep Dive & Differential Signaling"),
    ("Chapter 42", "Mathematical Proof of Covariance Modulated EKF"),
    ("Appendix A", "Glossary of Technical Terms"),
    ("Appendix B", "References (70 Verified Academic Citations)"),
    ("Appendix C", "Extended Bibliography (Further Reading)"),
    ("Appendix D", "Index of Key Mathematical Formulas"),
    ("Appendix E", "Quick-Reference Hardware Pinout Table")
]

# Insert entries in reverse so they appear in correct 1->42 order
for num, title in reversed(toc_entries):
    line_text = f"{num} — {title} "
    dots_needed = max(5, 100 - len(line_text))
    dots = "." * dots_needed
    full_toc_line = f"{line_text}{dots}"
    
    new_p = toc_heading_p.insert_paragraph_before(full_toc_line)
    new_p.paragraph_format.space_after = Pt(2)
    new_p.paragraph_format.space_before = Pt(1)
    new_p.paragraph_format.left_indent = Inches(0.2)
    
    run = new_p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6C)

doc.save(file_path)
print("SUCCESSFULLY RE-ORDERED TOC FROM CHAPTER 1 TO APPENDIX E!")
