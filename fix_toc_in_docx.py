import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

file_path = r"c:\Users\mksin\Desktop\AI hardened BMS\Cyber_Hardened_BMS_Manual.docx"
doc = docx.Document(file_path)

# First clean up any corrupted characters (like ? or  in headings)
for p in doc.paragraphs:
    if p.text.startswith("Chapter") or p.text.startswith("Appendix") or p.text.startswith("APPENDIX"):
        # Replace odd replacement chars with proper em-dash '—'
        new_text = p.text.replace(" ? ", " — ").replace("  ", " — ")
        if new_text != p.text:
            p.text = new_text

# Define full Table of Contents entries
toc_entries = [
    ("Chapter 1", "Executive Summary"),
    ("Chapter 2", "Introduction & Motivation"),
    ("Chapter 3", "Literature Review"),
    ("Chapter 4", "System Overview & Master Schematic Architecture"),
    ("Chapter 5", "Core Theory: How Everything Works"),
    ("Chapter 6", "Software & Tools Setup Guide"),
    ("Chapter 7", "Hardware Architecture & Bill of Materials"),
    ("Chapter 8", "Simulation Phase (LTspice & MATLAB/Simulink)"),
    ("Chapter 9", "Hardware Assembly (Safe, Step-by-Step)"),
    ("Chapter 10", "Firmware Development & Dual-Core Task Architecture"),
    ("Chapter 11", "Attack Bench & Dataset Generation"),
    ("Chapter 12", "ML Classifier Training & Deployment"),
    ("Chapter 13", "The IDS–EKF Feedback Loop (Patent Core)"),
    ("Chapter 14", "Testing, Validation & Deliverables Checklist"),
    ("Chapter 15", "12-Week Timeline & Milestones"),
    ("Chapter 16", "Writing the IEEE Conference Paper"),
    ("Chapter 17", "Patent Filing Guide (Indian Patent Office)"),
    ("Chapter 18", "Presenting to Your Professor & Viva Walkthrough"),
    ("Chapter 19", "Error Log: Every Correction Applied in This Manual"),
    ("Chapter 20", "Extended Kalman Filter: Full Mathematical Derivation"),
    ("Chapter 21", "Worked Numerical Examples & Step-by-Step Calculations"),
    ("Chapter 22", "Complete Firmware Source Code Listings"),
    ("Chapter 23", "Complete Python ML Pipeline Source Code"),
    ("Chapter 24", "Hardware & Software Troubleshooting Guide"),
    ("Chapter 25", "Frequently Asked Questions (FAQ)"),
    ("Chapter 26", "Team Roles & Daily Task Breakdown"),
    ("Chapter 27", "IEEE Paper: Full Draft Template"),
    ("Chapter 28", "Safety & Compliance Notes"),
    ("Chapter 29", "Component Datasheet Quick-Reference"),
    ("Chapter 30", "Standard Project Report Structure"),
    ("Chapter 31", "Patent Forms: Field-by-Field Guide"),
    ("Chapter 32", "Viva / Interview Questions & Model Answers"),
    ("Chapter 33", "Expected Results: Graph-by-Graph Description"),
    ("Chapter 34", "PCB Design Guide (KiCad Detail)"),
    ("Chapter 35", "Alternatives Considered and Rejected"),
    ("Chapter 36", "Data Logging & Post-Processing"),
    ("Chapter 37", "Environmental & Sustainability Notes"),
    ("Chapter 38", "Sample Data & Test Log Format"),
    ("Chapter 39", "Deliverables Mapped to Typical Evaluation Criteria"),
    ("Chapter 40", "Advanced Battery Diagnostics & Degradation Physics"),
    ("Chapter 41", "CAN Protocol Deep Dive & Differential Signaling"),
    ("Chapter 42", "Mathematical Proof of Covariance Modulated EKF"),
    ("Appendix A", "Glossary of Technical Terms"),
    ("Appendix B", "References (70 Verified Academic Citations)"),
    ("Appendix C", "Extended Bibliography (Further Reading)"),
    ("Appendix D", "Index of Key Mathematical Formulas"),
    ("Appendix E", "Quick-Reference Hardware Pinout Table")
]

# Find paragraph 8 ("TABLE OF CONTENTS")
toc_p_idx = -1
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "TABLE OF CONTENTS":
        toc_p_idx = i
        break

print(f"Found TABLE OF CONTENTS at paragraph index {toc_p_idx}")

if toc_p_idx != -1:
    toc_heading_p = doc.paragraphs[toc_p_idx]
    
    # Insert formatted TOC entries right after TABLE OF CONTENTS heading
    current_p = toc_heading_p
    for num, title in toc_entries:
        line_text = f"{num} — {title} "
        dots_needed = max(5, 110 - len(line_text))
        dots = "." * dots_needed
        full_toc_line = f"{line_text}{dots}"
        
        new_p = current_p.insert_paragraph_before(full_toc_line)
        new_p.paragraph_format.space_after = Pt(2)
        new_p.paragraph_format.space_before = Pt(1)
        new_p.paragraph_format.left_indent = Inches(0.2)
        
        run = new_p.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6C)
        
        current_p = new_p

doc.save(file_path)
print("SUCCESSFULLY FIXED TABLE OF CONTENTS IN DOCX!")
